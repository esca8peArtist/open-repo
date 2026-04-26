"""
Document ingestion pipeline.
Supports: PDF, DOCX, TXT, Markdown, web pages.
All processing is local — no cloud OCR or parsing.

Documents are chunked into overlapping segments, embedded via nomic-embed-text
through Ollama, and stored in ChromaDB under a per-agent collection.
Ingestion records are also written to the PostgreSQL documents table.
"""

import asyncio
import hashlib
import logging
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import List
from urllib.parse import urlparse

import httpx
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from agentcore.config import get_settings
from agentcore.server import get_chroma_client
from .embeddings import embed_batch

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# PostgreSQL helper (asyncpg — optional dependency at call sites)
# ---------------------------------------------------------------------------


async def _record_document_in_postgres(
    agent_id: str,
    source: str,
    file_type: str,
    chunk_count: int,
    metadata: dict,
) -> None:
    """Write a row to the documents table in PostgreSQL.

    Uses asyncpg directly. Failure is logged but does not abort ingestion
    so that the ChromaDB state remains the source of truth for RAG.
    """
    try:
        import asyncpg  # optional; available inside agentcore container

        conn = await asyncpg.connect(dsn=get_settings().postgres_url)
        try:
            await conn.execute(
                """
                INSERT INTO documents (agent_id, filename, file_type, chunk_count, metadata)
                VALUES ($1::uuid, $2, $3, $4, $5::jsonb)
                """,
                agent_id,
                source,
                file_type,
                chunk_count,
                __import__("json").dumps(metadata),
            )
        finally:
            await conn.close()
    except Exception as exc:
        logger.warning("Failed to record document in PostgreSQL: %s", exc)


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------


@dataclass
class DocumentChunk:
    text: str
    metadata: dict = field(default_factory=dict)
    # metadata keys: source, page, chunk_index, doc_id, file_type


# ---------------------------------------------------------------------------
# Text chunking
# ---------------------------------------------------------------------------


async def chunk_text(
    text: str, chunk_size: int = 512, overlap: int = 50
) -> List[str]:
    """Split text into overlapping chunks for embedding.

    Uses a word-boundary-aware sliding window. Words are never split in the
    middle. The chunk_size parameter is measured in characters (not tokens),
    which is a reasonable proxy for nomic-embed-text's 8192-token limit.

    Args:
        text:       Input text (may contain newlines).
        chunk_size: Target chunk size in characters. Default 512.
        overlap:    Number of characters of overlap between adjacent chunks.
                    Helps preserve context across chunk boundaries.

    Returns:
        List of text chunks. Each chunk is at most chunk_size characters
        (possibly a little more to avoid splitting in the middle of a word).
    """
    if not text or not text.strip():
        return []

    # Normalise whitespace: collapse multiple blank lines, strip edges.
    import re
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    if len(text) <= chunk_size:
        return [text]

    chunks: List[str] = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        if end >= len(text):
            # Last chunk — take everything remaining.
            chunks.append(text[start:].strip())
            break

        # Walk back from end to the nearest whitespace boundary so we don't
        # cut in the middle of a word.
        boundary = end
        while boundary > start and not text[boundary].isspace():
            boundary -= 1

        if boundary == start:
            # No whitespace found — hard cut at chunk_size.
            boundary = end

        chunk = text[start:boundary].strip()
        if chunk:
            chunks.append(chunk)

        # Advance start with overlap.
        start = boundary - overlap
        if start < 0:
            start = 0
        # Skip past leading whitespace after backing up by overlap.
        while start < len(text) and text[start].isspace():
            start += 1

    return [c for c in chunks if c]


# ---------------------------------------------------------------------------
# Format-specific extractors
# ---------------------------------------------------------------------------


async def ingest_pdf(path: Path, agent_id: str) -> List[DocumentChunk]:
    """Extract text from PDF using pypdf. Chunk by paragraph/page.

    Args:
        path:     Absolute path to the PDF file.
        agent_id: Agent identifier; used for metadata and ChromaDB collection.

    Returns:
        List of DocumentChunk objects ready for embedding and storage.
    """
    doc_id = str(uuid.uuid4())
    chunks: List[DocumentChunk] = []

    def _read_pdf() -> List[tuple[int, str]]:
        """Synchronous PDF reading; run in thread executor."""
        reader = PdfReader(str(path))
        pages: List[tuple[int, str]] = []
        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append((page_num, page_text))
        return pages

    # pypdf is CPU-bound; offload to thread pool so we don't block the event loop.
    loop = asyncio.get_running_loop()
    pages = await loop.run_in_executor(None, _read_pdf)

    chunk_idx = 0
    for page_num, page_text in pages:
        page_chunks = await chunk_text(page_text)
        for chunk in page_chunks:
            chunks.append(
                DocumentChunk(
                    text=chunk,
                    metadata={
                        "source": str(path),
                        "page": page_num,
                        "chunk_index": chunk_idx,
                        "doc_id": doc_id,
                        "file_type": "pdf",
                        "agent_id": agent_id,
                    },
                )
            )
            chunk_idx += 1

    logger.info(
        "PDF ingestion: %s — %d pages → %d chunks",
        path.name,
        len(pages),
        len(chunks),
    )
    return chunks


async def ingest_docx(path: Path, agent_id: str) -> List[DocumentChunk]:
    """Extract text from DOCX using python-docx.

    Extracts paragraph text in document order. Tables are extracted as
    tab-separated rows. Chunks at paragraph boundaries.

    Args:
        path:     Absolute path to the DOCX file.
        agent_id: Agent identifier.

    Returns:
        List of DocumentChunk objects.
    """
    doc_id = str(uuid.uuid4())

    def _read_docx() -> str:
        """Synchronous DOCX reading; run in thread executor."""
        doc = DocxDocument(str(path))
        parts: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        return "\n\n".join(parts)

    loop = asyncio.get_running_loop()
    full_text = await loop.run_in_executor(None, _read_docx)

    raw_chunks = await chunk_text(full_text)
    chunks = [
        DocumentChunk(
            text=chunk,
            metadata={
                "source": str(path),
                "page": None,
                "chunk_index": idx,
                "doc_id": doc_id,
                "file_type": "docx",
                "agent_id": agent_id,
            },
        )
        for idx, chunk in enumerate(raw_chunks)
    ]

    logger.info("DOCX ingestion: %s → %d chunks", path.name, len(chunks))
    return chunks


async def ingest_text(path: Path, agent_id: str) -> List[DocumentChunk]:
    """Read plain text or Markdown. Chunk by paragraphs.

    Args:
        path:     Absolute path to a .txt or .md file.
        agent_id: Agent identifier.

    Returns:
        List of DocumentChunk objects.
    """
    doc_id = str(uuid.uuid4())

    def _read_file() -> str:
        return path.read_text(encoding="utf-8", errors="replace")

    loop = asyncio.get_running_loop()
    content = await loop.run_in_executor(None, _read_file)

    raw_chunks = await chunk_text(content)
    suffix = path.suffix.lstrip(".").lower() or "txt"
    file_type = "md" if suffix == "md" else "txt"

    chunks = [
        DocumentChunk(
            text=chunk,
            metadata={
                "source": str(path),
                "page": None,
                "chunk_index": idx,
                "doc_id": doc_id,
                "file_type": file_type,
                "agent_id": agent_id,
            },
        )
        for idx, chunk in enumerate(raw_chunks)
    ]

    logger.info("Text ingestion: %s → %d chunks", path.name, len(chunks))
    return chunks


async def ingest_url(url: str, agent_id: str) -> List[DocumentChunk]:
    """Scrape web page with httpx + BeautifulSoup. Extract main content.

    Fetches the URL, parses HTML, and extracts meaningful text from the
    <main>, <article>, or <body> element. Scripts, styles, nav, footer,
    and header elements are stripped.

    Args:
        url:      Full HTTP/HTTPS URL to scrape.
        agent_id: Agent identifier.

    Returns:
        List of DocumentChunk objects.

    Raises:
        httpx.HTTPStatusError: If the URL returns a non-2xx response.
    """
    doc_id = hashlib.sha256(url.encode()).hexdigest()[:32]

    async with httpx.AsyncClient(
        timeout=30.0,
        follow_redirects=True,
        headers={"User-Agent": "AgentCore/1.0 (+https://agent.local)"},
    ) as client:
        response = await client.get(url)
        response.raise_for_status()
        html = response.text

    soup = BeautifulSoup(html, "html.parser")

    # Remove boilerplate elements.
    for tag in soup.find_all(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    # Prefer semantic content elements; fall back to full body.
    content_el = (
        soup.find("main")
        or soup.find("article")
        or soup.find(id="content")
        or soup.find(class_="content")
        or soup.find("body")
        or soup
    )

    # Extract text with meaningful whitespace.
    text = content_el.get_text(separator="\n", strip=True)

    raw_chunks = await chunk_text(text)
    domain = urlparse(url).netloc

    chunks = [
        DocumentChunk(
            text=chunk,
            metadata={
                "source": url,
                "domain": domain,
                "page": None,
                "chunk_index": idx,
                "doc_id": doc_id,
                "file_type": "url",
                "agent_id": agent_id,
            },
        )
        for idx, chunk in enumerate(raw_chunks)
    ]

    logger.info("URL ingestion: %s → %d chunks", url, len(chunks))
    return chunks


# ---------------------------------------------------------------------------
# Storage helper
# ---------------------------------------------------------------------------


async def _store_chunks_in_chroma(
    chunks: List[DocumentChunk], agent_id: str
) -> None:
    """Embed all chunks and upsert them into the agent's ChromaDB collection.

    Collection name: f"agent_{agent_id}" (ChromaDB collection per agent).
    IDs are deterministic hashes of (doc_id, chunk_index) to make re-ingestion
    idempotent — same document ingested twice overwrites existing chunks.
    """
    if not chunks:
        return

    client = get_chroma_client()
    collection_name = f"agent_{agent_id}"

    collection = await client.get_or_create_collection(
        name=collection_name,
        metadata={"hnsw:space": "cosine"},  # cosine similarity for nomic-embed-text
    )

    texts = [c.text for c in chunks]
    metadatas = [c.metadata for c in chunks]

    # Deterministic IDs: sha256(doc_id + chunk_index) → reproducible upserts.
    ids = [
        hashlib.sha256(
            f"{c.metadata.get('doc_id', '')}_{c.metadata.get('chunk_index', i)}".encode()
        ).hexdigest()[:36]
        for i, c in enumerate(chunks)
    ]

    # Generate embeddings in batch (handles concurrency internally).
    embeddings = await embed_batch(texts)

    await collection.upsert(
        ids=ids,
        documents=texts,
        embeddings=embeddings,
        metadatas=metadatas,
    )

    logger.info(
        "Stored %d chunks in ChromaDB collection '%s'",
        len(chunks),
        collection_name,
    )


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------


async def ingest_document(source: str | Path, agent_id: str) -> int:
    """Auto-detect type and ingest. Returns number of chunks stored.

    Detects the document type from the file extension (for local files) or
    from the URL scheme (for web pages). Delegates to the appropriate
    ingest_* function, then stores chunks + embeddings in ChromaDB and
    records the document in the PostgreSQL documents table.

    Args:
        source:   File path (str or Path) or HTTP/HTTPS URL string.
        agent_id: The agent this document belongs to. Determines which
                  ChromaDB collection the chunks are stored in.

    Returns:
        Number of chunks stored in ChromaDB.

    Raises:
        ValueError: If the source type cannot be determined.
        FileNotFoundError: If the source is a local path that does not exist.
    """
    source_str = str(source)
    chunks: List[DocumentChunk] = []
    file_type: str = "unknown"

    # Determine source type.
    if source_str.startswith("http://") or source_str.startswith("https://"):
        chunks = await ingest_url(source_str, agent_id)
        file_type = "url"
    else:
        path = Path(source_str)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        suffix = path.suffix.lower()
        if suffix == ".pdf":
            chunks = await ingest_pdf(path, agent_id)
            file_type = "pdf"
        elif suffix in (".docx", ".doc"):
            chunks = await ingest_docx(path, agent_id)
            file_type = "docx"
        elif suffix in (".txt", ".md", ".markdown", ".rst"):
            chunks = await ingest_text(path, agent_id)
            file_type = "txt" if suffix == ".txt" else "md"
        else:
            raise ValueError(
                f"Unsupported document type: {suffix}. "
                "Supported: .pdf, .docx, .txt, .md, or an HTTP/HTTPS URL."
            )

    if not chunks:
        logger.warning("No chunks produced from source: %s", source_str)
        return 0

    # Store embeddings in ChromaDB.
    await _store_chunks_in_chroma(chunks, agent_id)

    # Record in PostgreSQL (best-effort).
    await _record_document_in_postgres(
        agent_id=agent_id,
        source=source_str,
        file_type=file_type,
        chunk_count=len(chunks),
        metadata={
            "doc_id": chunks[0].metadata.get("doc_id", ""),
            "source": source_str,
        },
    )

    logger.info(
        "Ingestion complete: %s → %d chunks (agent=%s)",
        source_str,
        len(chunks),
        agent_id,
    )
    return len(chunks)
